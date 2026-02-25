from fastapi import FastAPI, UploadFile, File, Depends, HTTPException, status, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.responses import FileResponse
from dotenv import load_dotenv
from groq import Groq
import os
import json
import shutil
import sqlite3
from datetime import timedelta
from typing import Optional
import asyncio
from docx import Document
from collections import defaultdict

# import logging
# Update this line (around line 17)
from rag_docling import process_document, search_chunks, detect_format, collection, get_chunks_by_doc, get_tables_by_doc
from auth import (
    init_db, create_user, get_user, verify_password, create_access_token,
    get_current_user, Token, ACCESS_TOKEN_EXPIRE_MINUTES, create_document,
    get_user_documents, get_document_owner, DB_NAME
)
from web_search import process_web_search, find_relevant_chunks
from agents.orchestrator import Orchestrator
from agents.master_agent import MasterAgent
from agents.compliance_agent import ComplianceAgent
from agents.training_agent import TrainingAgent
from chat_history import ChatHistoryManager
from gtts import gTTS
import tempfile
import base64
import uuid

load_dotenv()

# Initialize User DB
init_db()

# Initialize Chat History Manager
chat_history_manager = ChatHistoryManager()

# Global Progress Store
# upload_progress = {}

# def update_progress(doc_id, status, percentage):
#     upload_progress[doc_id] = {"status": status, "percentage": percentage}


# Use Groq client
client = Groq(api_key=os.getenv("GROQ_API_KEY"))
print(os.getenv("GROQ_API_KEY"))
 
app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

LLM_MODEL = "openai/gpt-oss-120b"
LLM_VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"  # used in vision_figures

UPLOAD_DIR = "uploads"
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

@app.post("/register")
async def register(form_data: OAuth2PasswordRequestForm = Depends()):
    if create_user(form_data.username, form_data.password):
        return {"message": "User created successfully"}
    else:
        raise HTTPException(status_code=400, detail="Username already registered")

@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    user = get_user(form_data.username)
    if not user or not verify_password(form_data.password, user["hashed_password"]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user["username"]}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/ingest")
async def ingest(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """
    Ingest documents using Docling.
    
    Supports multiple formats:
    - PDF (with OCR for scanned documents)
    - Word (DOCX)
    - PowerPoint (PPTX)
    - Excel (XLSX)
    - Images (PNG, JPEG, TIFF, etc.)
    - HTML
    
    Features:
    - Advanced layout understanding
    - Table structure extraction
    - Image classification
    - Intelligent chunking
    - Rich metadata extraction
    """
    try:
        # Detect format
        file_format = detect_format(file.filename)
        
        # Create document record
        doc_id = create_document(current_user["id"], file.filename)
        
        # Save file to disk (required for Docling)
        file_path = os.path.join(UPLOAD_DIR, f"{doc_id}_{file.filename}")
        with open(file_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
        
        # Process with Docling using file path
        stats = process_document(
            file_path=file_path,  # Pass file path instead of bytes
            filename=file.filename,
            user_id=current_user["id"],
            doc_id=doc_id
        )
        
        return {
            "status": "success",
            "message": f"Document processed successfully using Docling",
            "doc_id": doc_id,
            "filename": file.filename,
            "format": file_format,
            "stats": stats
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error processing document: {str(e)}"
        )

@app.get("/documents/{doc_id}/file")
async def get_document_file(doc_id: int, current_user: dict = Depends(get_current_user)):
    """Download original document file"""
    owner_id = get_document_owner(doc_id)
    if owner_id != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized")
    
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT filename FROM documents WHERE id = ?", (doc_id,))
    row = c.fetchone()
    conn.close()
    
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    
    filename = row[0]
    file_path = os.path.join(UPLOAD_DIR, f"{doc_id}_{filename}")
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found on server")
    
    # Detect media type
    ext = filename.lower().split('.')[-1]
    media_types = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'png': 'image/png',
        'jpg': 'image/jpeg',
        'jpeg': 'image/jpeg',
    }
    media_type = media_types.get(ext, 'application/octet-stream')
    
    return FileResponse(file_path, media_type=media_type, filename=filename)

@app.get("/documents")
async def get_documents(current_user: dict = Depends(get_current_user)):
    """List all documents for current user"""
    return get_user_documents(current_user["id"])

@app.post("/query")
async def query(question: str, doc_id: Optional[int] = None, current_user: dict = Depends(get_current_user)):
    """
    Query documents using RAG with Docling-processed content.
    
    Enhanced with:
    - Richer context from better document understanding
    - Metadata like page numbers, sections, and content types
    - Better handling of tables, images, and complex layouts
    """
    # Verify document ownership if doc_id provided
    if doc_id:
        owner_id = get_document_owner(doc_id)
        if owner_id != current_user["id"]:
            raise HTTPException(status_code=403, detail="Not authorized to access this document")
    
    # Search for relevant chunks
    top_chunks = search_chunks(
        query=question,
        user_id=current_user["id"],
        doc_id=doc_id,
        k=8 if doc_id else 5,
    )
    
    if not top_chunks:
        return {
            "response": {
                "answer": "No relevant information found in the selected document(s).",
                "citations": []
            },
            "chunks_used": []
        }
    
    # Build context with enhanced metadata
    context_parts = []
    for chunk in top_chunks:
        section_info = f" [{chunk['section']}]" if chunk['section'] else ""
        page_info = f"Page {chunk['page']}"
        context_parts.append(f"{page_info}{section_info}:\n{chunk['text']}")
    
    context = "\n\n---\n\n".join(context_parts)
    
    # Enhanced prompt with metadata awareness and markdown formatting
    prompt = f"""You are a RAG assistant with access to document content that has been carefully extracted and structured.

The context below includes page numbers and section headings for precise citations.

Context:
{context}

Question: {question}

Instructions:
1. Answer the question based ONLY on the provided context
2. If the answer is not in the context, say "Not found in document"
3. Include specific page numbers and sections when citing information
4. For tables or structured data, preserve the structure in your answer
5. Format your answer using Markdown syntax for better readability:
   - Use **bold** for emphasis
   - Use *italic* for subtle emphasis
   - Use bullet points (- or *) for lists
   - Use numbered lists (1., 2., etc.) for ordered items
   - Use code blocks (```language) for code snippets
   - Use `inline code` for technical terms
   - Use ## for section headings if the answer is long
   - Use > for blockquotes when citing specific passages
   - Use tables when presenting structured data

IMPORTANT: Return your answer in Markdown format. Make it well-structured, easy to read, and visually appealing.

Return your response as JSON in this exact format:
{{
  "answer": "your detailed markdown-formatted answer here",
  "citations": [
    {{"page": 1, "section": "Introduction", "snippet": "relevant text"}},
    {{"page": 3, "section": "Methods", "snippet": "relevant text"}}
  ]
}}
"""
    
    try:
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}]
        )
        
        raw_output = response.choices[0].message.content
        
        # Parse JSON response
        try:
            data = json.loads(raw_output)
        except:
            # Clean up markdown code blocks
            cleaned = raw_output.replace("```json", "").replace("```", "").strip()
            data = json.loads(cleaned)
        
        return {
            "response": data,
            "chunks_used": top_chunks,
            "metadata": {
                "model": LLM_MODEL,
                "chunks_retrieved": len(top_chunks),
                "document_specific": doc_id is not None
            }
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error generating response: {str(e)}"
        )

# 


        

@app.post("/web-query")
async def web_query(question: str, current_user: dict = Depends(get_current_user)):
    """
    Query the web using DuckDuckGo search with intelligent scraping.
    Returns structured response with citations similar to Perplexity.
    
    Features:
    - Free web search via DuckDuckGo (no API key required)
    - Reliable web scraping with quality filtering
    - Semantic search for relevant content
    - Structured citations with URLs
    """
    try:
        # Step 1: Search web and extract content
        web_data = await process_web_search(query=question, num_results=6)
        
        if 'error' in web_data or not web_data['chunks']:
            return {
                "response": {
                    "answer": "I couldn't find relevant information on the web for your query. This might be because:\n- The search didn't return accessible results\n- Websites blocked scraping\n- The query might be too specific or misspelled\n\nPlease try rephrasing your question.",
                    "citations": []
                },
                "sources_used": 0,
                "metadata": {
                    "error": web_data.get('error', 'Unknown error')
                }
            }
        
        # Step 2: Find most relevant chunks
        relevant_chunks = find_relevant_chunks(
            query=question,
            chunks=web_data['chunks'],
            sources=web_data['sources'],
            embeddings=web_data['embeddings'],
            k=8  # Get more chunks for better context
        )
        
        if not relevant_chunks:
            return {
                "response": {
                    "answer": "No relevant information found in the scraped content.",
                    "citations": []
                },
                "sources_used": 0
            }
        
        # Step 3: Build context for LLM
        context_parts = []
        sources_map = {}
        
        for i, chunk_data in enumerate(relevant_chunks):
            source = chunk_data['source']
            url = source['url']
            
            # Track unique sources
            if url not in sources_map:
                snippet = source.get('snippet') or chunk_data['text'][:200]
                sources_map[url] = {
                    'url': url,
                    'title': source['title'],
                    'domain': source['domain'],
                    'snippet': snippet
                }
            
            # Add context with source reference
            context_parts.append(
                f"[Source {i+1}] {source['title']} ({source['domain']}):\n{chunk_data['text']}\n"
            )
        
        context = "\n---\n\n".join(context_parts)
        
        # Step 4: Generate response with LLM
        prompt = f"""You are a helpful AI search assistant. Answer the user's question based on the web search results provided below.

USER QUESTION:
{question}

WEB SEARCH RESULTS:
{context}

INSTRUCTIONS:
1. Provide a comprehensive, well-structured answer based on the search results
2. Synthesize information from multiple sources when possible
3. Be accurate and cite your sources naturally in the text
4. If the search results don't fully answer the question, acknowledge this
5. Use clear, concise language
6. Format your answer using Markdown syntax for better readability:
   - Use **bold** for emphasis and key points
   - Use *italic* for subtle emphasis
   - Use bullet points (- or *) for lists
   - Use numbered lists (1., 2., etc.) for ordered items
   - Use code blocks (```language) for code snippets or technical examples
   - Use `inline code` for technical terms, product names, or specific values
   - Use ## for section headings to organize longer answers
   - Use > for blockquotes when citing specific information
   - Use tables when presenting structured data or comparisons
   - Use horizontal rules (---) to separate major sections

IMPORTANT: Return your answer in Markdown format. Make it well-structured, visually appealing, and easy to read.

Return ONLY a JSON object with this EXACT structure:
{{
  "answer": "Your detailed markdown-formatted answer here. Reference sources like 'According to [source title]...' naturally in your text.",
  "key_sources": [
    {{"url": "https://example.com", "title": "Page Title", "relevance": "Brief note on what this source contributed"}}
  ]
}}"""
        
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3  # Lower temperature for more factual responses
        )
        
        raw_output = response.choices[0].message.content.strip()
        
        # Parse JSON response
        try:
            # Remove markdown code blocks if present
            if raw_output.startswith('```'):
                raw_output = raw_output.split('```')[1]
                if raw_output.startswith('json'):
                    raw_output = raw_output[4:]
                raw_output = raw_output.strip()
            
            data = json.loads(raw_output)
        except json.JSONDecodeError as e:
            print(f"JSON parse error: {e}\nRaw output: {raw_output}")
            # Fallback: return raw text
            return {
                "response": {
                    "answer": raw_output,
                    "citations": list(sources_map.values())[:3]
                },
                "sources_used": len(sources_map),
                "metadata": {
                    "warning": "LLM returned non-JSON response"
                }
            }
        
        # Enrich citations
        key_sources = data.get('key_sources', [])
        enriched_citations = []
        
        for citation in key_sources[:5]:  # Limit to top 5 sources
            url = citation.get('url', '')
            if url in sources_map:
                enriched_citations.append({
                    **sources_map[url],
                    'relevance': citation.get('relevance', '')
                })
        
        # If no key sources, use top 3 from sources_map
        if not enriched_citations:
            enriched_citations = list(sources_map.values())[:3]
        
        return {
            "response": {
                "answer": data.get('answer', ''),
                "citations": enriched_citations
            },
            "sources_used": len(sources_map),
            "chunks_analyzed": len(relevant_chunks),
            "metadata": {
                "model": LLM_MODEL,
                "search_engine": "DuckDuckGo",
            }
        }
        try:
            data = json.loads(raw_output)
        except:
            # Clean up markdown code blocks
            cleaned = raw_output.replace("```json", "").replace("```", "").strip()
            data = json.loads(cleaned)
        
        return {
            "response": data,
            "chunks_used": top_chunks,
            "metadata": {
                "model": LLM_MODEL,
                "chunks_retrieved": len(top_chunks),
                "document_specific": doc_id is not None
            }
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error generating response: {str(e)}"
        )





        

@app.post("/web-query")
async def web_query(question: str, current_user: dict = Depends(get_current_user)):
    """
    Query the web using DuckDuckGo search with intelligent scraping.
    Returns structured response with citations similar to Perplexity.
    
    Features:
    - Free web search via DuckDuckGo (no API key required)
    - Reliable web scraping with quality filtering
    - Semantic search for relevant content
    - Structured citations with URLs
    """
    try:
        # Step 1: Search web and extract content
        web_data = await process_web_search(query=question, num_results=6)
        
        if 'error' in web_data or not web_data['chunks']:
            return {
                "response": {
                    "answer": "I couldn't find relevant information on the web for your query. This might be because:\n- The search didn't return accessible results\n- Websites blocked scraping\n- The query might be too specific or misspelled\n\nPlease try rephrasing your question.",
                    "citations": []
                },
                "sources_used": 0,
                "metadata": {
                    "error": web_data.get('error', 'Unknown error')
                }
            }
        
        # Step 2: Find most relevant chunks
        relevant_chunks = find_relevant_chunks(
            query=question,
            chunks=web_data['chunks'],
            sources=web_data['sources'],
            embeddings=web_data['embeddings'],
            k=8  # Get more chunks for better context
        )
        
        if not relevant_chunks:
            return {
                "response": {
                    "answer": "No relevant information found in the scraped content.",
                    "citations": []
                },
                "sources_used": 0
            }
        
        # Step 3: Build context for LLM
        context_parts = []
        sources_map = {}
        
        for i, chunk_data in enumerate(relevant_chunks):
            source = chunk_data['source']
            url = source['url']
            
            # Track unique sources
            if url not in sources_map:
                snippet = source.get('snippet') or chunk_data['text'][:200]
                sources_map[url] = {
                    'url': url,
                    'title': source['title'],
                    'domain': source['domain'],
                    'snippet': snippet
                }
            
            # Add context with source reference
            context_parts.append(
                f"[Source {i+1}] {source['title']} ({source['domain']}):\n{chunk_data['text']}\n"
            )
        
        context = "\n---\n\n".join(context_parts)
        
        # Step 4: Generate response with LLM
        prompt = f"""You are a helpful AI search assistant. Answer the user's question based on the web search results provided below.

USER QUESTION:
{question}

WEB SEARCH RESULTS:
{context}

INSTRUCTIONS:
1. Provide a comprehensive, well-structured answer based on the search results
2. Synthesize information from multiple sources when possible
3. Be accurate and cite your sources naturally in the text
4. If the search results don't fully answer the question, acknowledge this
5. Use clear, concise language
6. Format your answer using Markdown syntax for better readability:
   - Use **bold** for emphasis and key points
   - Use *italic* for subtle emphasis
   - Use bullet points (- or *) for lists
   - Use numbered lists (1., 2., etc.) for ordered items
   - Use code blocks (```language) for code snippets or technical examples
   - Use `inline code` for technical terms, product names, or specific values
   - Use ## for section headings to organize longer answers
   - Use > for blockquotes when citing specific information
   - Use tables when presenting structured data or comparisons
   - Use horizontal rules (---) to separate major sections

IMPORTANT: Return your answer in Markdown format. Make it well-structured, visually appealing, and easy to read.

Return ONLY a JSON object with this EXACT structure:
{{
  "answer": "Your detailed markdown-formatted answer here. Reference sources like 'According to [source title]...' naturally in your text.",
  "key_sources": [
    {{"url": "https://example.com", "title": "Page Title", "relevance": "Brief note on what this source contributed"}}
  ]
}}"""
        
        response = client.chat.completions.create(
            model=LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3  # Lower temperature for more factual responses
        )
        
        raw_output = response.choices[0].message.content.strip()
        
        # Parse JSON response
        try:
            # Remove markdown code blocks if present
            if raw_output.startswith('```'):
                raw_output = raw_output.split('```')[1]
                if raw_output.startswith('json'):
                    raw_output = raw_output[4:]
                raw_output = raw_output.strip()
            
            data = json.loads(raw_output)
        except json.JSONDecodeError as e:
            print(f"JSON parse error: {e}\nRaw output: {raw_output}")
            # Fallback: return raw text
            return {
                "response": {
                    "answer": raw_output,
                    "citations": list(sources_map.values())[:3]
                },
                "sources_used": len(sources_map),
                "metadata": {
                    "warning": "LLM returned non-JSON response"
                }
            }
        
        # Enrich citations
        key_sources = data.get('key_sources', [])
        enriched_citations = []
        
        for citation in key_sources[:5]:  # Limit to top 5 sources
            url = citation.get('url', '')
            if url in sources_map:
                enriched_citations.append({
                    **sources_map[url],
                    'relevance': citation.get('relevance', '')
                })
        
        # If no key sources, use top 3 from sources_map
        if not enriched_citations:
            enriched_citations = list(sources_map.values())[:3]
        
        return {
            "response": {
                "answer": data.get('answer', ''),
                "citations": enriched_citations
            },
            "sources_used": len(sources_map),
            "chunks_analyzed": len(relevant_chunks),
            "metadata": {
                "model": LLM_MODEL,
                "search_engine": "DuckDuckGo",
                "total_sources_found": web_data.get('total_sources', 0),
                "successful_crawls": web_data.get('successful_crawls', 0)
            }
        }
        
    except Exception as e:
        print(f"Web query error: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Error processing web query: {str(e)}"
        )

@app.post("/diagnose")
async def diagnose(question: str, current_user: dict = Depends(get_current_user)):
    """
    Multi-Agent Diagnostic Endpoint.
    
    Orchestrates multiple specialized AI agents to provide a comprehensive diagnosis:
    1. Intent Detection (Orchestrator)
    2. Parallel Execution:
       - Symptom Extraction
       - Sensor Data Analysis (Real-time/Mock)
       - Document Retrieval (RAG)
       - Historical Issue Matching
    3. Synthesis (Master Agent)
    """
    try:
        # Initialize Orchestrator
        orchestrator = Orchestrator()
        
        # Step 1: Detect Intent
        intent = await orchestrator.detect_intent(question)
        
        if intent == "diagnostic":
            # Step 2: Run Master Diagnostic Agent
            master_agent = MasterAgent()
            diagnosis = await master_agent.diagnose(question, current_user["id"])
            
            return {
                "type": "diagnostic",
                "result": diagnosis
            }
            
        elif intent == "compliance":
            # Run Compliance Agent
            compliance_agent = ComplianceAgent()
            report = await compliance_agent.check_compliance(question, current_user["id"])
            
            return {
                "type": "compliance",
                "result": report
            }
            
        elif intent == "training":
            # Run Training Agent
            training_agent = TrainingAgent()
            module = await training_agent.generate_training(question, current_user["id"])
            
            return {
                "type": "training",
                "result": module
            }
            
        else:
            # Fallback to standard RAG or generic response
            return {
                "type": "general",
                "message": "Query classified as general. Please use the standard query endpoint for general questions."
            }
            
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error in diagnostic process: {str(e)}"
        )
# ----- Chat History Endpoints -----

@app.get("/history/sessions")
async def get_sessions(current_user: dict = Depends(get_current_user)):
    return chat_history_manager.get_sessions(str(current_user["id"]))

@app.post("/history/sessions")
async def create_session(title: str = "New Chat", current_user: dict = Depends(get_current_user)):
    session_id = chat_history_manager.create_session(str(current_user["id"]), title)
    return {"session_id": session_id}

@app.get("/history/sessions/{session_id}/messages")
async def get_messages(session_id: str, current_user: dict = Depends(get_current_user)):
    # In a real app, verify user owns session
    return chat_history_manager.get_messages(session_id)

@app.post("/history/sessions/{session_id}/messages")
async def add_message(session_id: str, message: dict, current_user: dict = Depends(get_current_user)):
    # message = {"role": "user", "content": "...", "citations": [...]}
    chat_history_manager.add_message(
        session_id, 
        message["role"], 
        message["content"], 
        message.get("citations")
    )
    return {"status": "success"}

@app.delete("/history/sessions/{session_id}")
async def delete_session(session_id: str, current_user: dict = Depends(get_current_user)):
    chat_history_manager.delete_session(session_id)
    return {"status": "success"}

# ----- Voice Endpoints -----

@app.post("/voice/transcribe")
async def transcribe_audio(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    try:
        # Save temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name
        
        # Use Groq for transcription
        with open(tmp_path, "rb") as audio_file:
            transcription = client.audio.transcriptions.create(
                file=(tmp_path, audio_file.read()),
                model="distil-whisper-large-v3-en",
                response_format="json",
                language="en",
                temperature=0.0
            )
        
        os.unlink(tmp_path)
        return {"text": transcription.text}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    



@app.post("/voice/speak")
async def text_to_speech(text: str = Form(...), current_user: dict = Depends(get_current_user)):
    try:
        # Use gTTS
        tts = gTTS(text=text, lang='en')
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
            tts.save(tmp.name)
            tmp_path = tmp.name
            
        return FileResponse(tmp_path, media_type="audio/mpeg", filename="speech.mp3")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ----- Debug Endpoints -----

@app.get("/debug/vector-db")
async def debug_db(current_user: dict = Depends(get_current_user)):
    """
    Check the status of the ChromaDB collection.
    Requires authentication.
    """
    try:
        # returns the total number of items in your vector database
        count = collection.count()
        return {
            "status": "connected",
            "collection_name": "docling_chunks",
            "total_stored_chunks": count,
            "user_id": current_user["id"]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Vector DB Error: {str(e)}")
    
# ----- Progress Endpoint -----



@app.post("/generate-report")
async def generate_report(
    doc_id: int,
    current_user: dict = Depends(get_current_user)
):
    # -----------------------------
    # 1. AUTH CHECK
    # -----------------------------
    owner_id = get_document_owner(doc_id)
    if owner_id != current_user["id"]:
        raise HTTPException(status_code=403, detail="Not authorized")

    # -----------------------------
    # 2. FETCH DOC CHUNKS
    # -----------------------------
    chunks = get_chunks_by_doc(
        doc_id=doc_id,
        user_id=current_user["id"]
    )

    if not chunks:
        raise HTTPException(status_code=404, detail="No document content found in database.")

    # -----------------------------
    # 3. TAKE CONTENT SAFELY (WITH FALLBACK)
    # -----------------------------
    selected_text = []
    selected_pages = set()

    for c in chunks:
        # Safely get page number. Default to 1 if it's an image/has no page metadata.
        # Allow 0, 1, and 2 in case the parser uses 0-based indexing.
        page_num = c.get("page", 1)
        if page_num in [0, 1, 2]:
            selected_text.append(c.get("text", ""))
            selected_pages.add(page_num)

    # FALLBACK: If strict page filtering found nothing, just take the first 15 chunks
    if not selected_text:
        for c in chunks[:15]:
            selected_text.append(c.get("text", ""))
            selected_pages.add(c.get("page", 1))

    if not selected_text:
        raise HTTPException(status_code=404, detail="No extractable text found to generate report.")

    combined_text = "\n\n".join(selected_text)

    # -----------------------------
    # 4. FETCH RAW TABLES SAFELY
    # -----------------------------
    raw_tables = get_tables_by_doc(
        doc_id=doc_id,
        user_id=current_user["id"]
    )

    # Filter tables for pages 0, 1, 2
    page_filtered_tables = [
        t for t in raw_tables
        if t.get("page", 1) in [0, 1, 2]
    ]
    
    # Fallback: if no tables match the page filter, but tables DO exist, take the first 3
    if not page_filtered_tables and raw_tables:
        page_filtered_tables = raw_tables[:3]

    table_summary_lines = []
    for t_idx, t in enumerate(page_filtered_tables):
        table_summary_lines.append(f"\n[TABLE {t_idx + 1} — Page {t.get('page', 1)}]")
        for row in t.get("rows", []):
            table_summary_lines.append(" | ".join(str(cell).strip() for cell in row))

    table_summary_text = "\n".join(table_summary_lines) if table_summary_lines else ""

    full_content_for_llm = combined_text
    if table_summary_text:
        full_content_for_llm += "\n\n--- STRUCTURED TABLES FROM DOCUMENT ---\n" + table_summary_text

    # -----------------------------
    # 5. LLM – METALLURGICAL EXTRACTION
    # -----------------------------
    prompt = f"""
You are a senior metallurgical engineer.

Extract structured metallurgical information
from the given document content.

The content below includes both paragraph text and structured tables
extracted directly from the source document.
Prefer table data for chemical composition and mechanical properties
since it is more precise than narrative text.

STRICT RULES:
- No assumptions
- Preserve units, limits, standards
- If missing, write "Not specified"

Return JSON ONLY:

{{
  "header": {{
    "material": "",
    "standard": "",
    "doc_no": "CRM/MS/01"
  }},
  "scope": "",
  "specifications": {{
    "code": "",
    "colour": "",
    "equivalent_specs": [],
    "standards": []
  }},
  "chemical_composition": {{
    "min": {{}},
    "max": {{}}
  }},
  "mechanical_properties": [
    {{
      "property": "",
      "min": "",
      "max": "",
      "unit": ""
    }}
  ],
  "metallurgical_properties": {{
    "grain_size": "",
    "inclusion_rating": "",
    "microstructure": "",
    "surface_condition": "",
    "heat_treatment": ""
  }}
}}

SOURCE PAGES: {sorted(list(selected_pages))}

CONTENT:
{full_content_for_llm}
"""

    response = await asyncio.to_thread(
        client.chat.completions.create,
        model=LLM_MODEL, # Ensure this variable is defined globally
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        response_format={"type": "json_object"}
    )

    data = json.loads(response.choices[0].message.content)

    # -----------------------------
    # 6. CREATE METALLURGY DOCX
    # -----------------------------
    os.makedirs("reports", exist_ok=True)
    report_path = f"reports/doc_{doc_id}_metallurgy_report.docx"

    doc = Document()

    # -------- HEADER TABLE --------
    header = doc.add_table(rows=2, cols=4)
    header.style = "Table Grid"

    header.cell(0, 1).text = "Materials & Metallurgy Group"
    header.cell(0, 3).text = data.get("header", {}).get("doc_no", "N/A")

    header.cell(1, 1).text = "Material Specification"
    header.cell(1, 2).text = data.get("header", {}).get("material", "N/A")

    doc.add_paragraph("")

    # -------- 1. SCOPE --------
    doc.add_heading("1. SCOPE", level=2)
    doc.add_paragraph(data.get("scope", "Not specified"))

    # -------- 2. SPECIFICATIONS --------
    doc.add_heading("2. SPECIFICATIONS", level=2)
    spec_table = doc.add_table(rows=4, cols=2)
    spec_table.style = "Table Grid"

    specs = data.get("specifications", {})
    spec_table.cell(0, 0).text = "Code"
    spec_table.cell(0, 1).text = specs.get("code", "-")

    spec_table.cell(1, 0).text = "Colour"
    spec_table.cell(1, 1).text = specs.get("colour", "-")

    spec_table.cell(2, 0).text = "Equivalent Specs"
    spec_table.cell(2, 1).text = ", ".join(specs.get("equivalent_specs", []))

    spec_table.cell(3, 0).text = "Standards"
    spec_table.cell(3, 1).text = ", ".join(specs.get("standards", []))

    # -------- 3. CHEMICAL COMPOSITION --------
    doc.add_heading("3. CHEMICAL COMPOSITION", level=2)

    chem = data.get("chemical_composition", {"min": {}, "max": {}})
    elements = list(chem.get("min", {}).keys())

    if elements:
        chem_table = doc.add_table(rows=3, cols=len(elements) + 1)
        chem_table.style = "Table Grid"

        chem_table.cell(0, 0).text = "Element"
        for i, el in enumerate(elements):
            chem_table.cell(0, i + 1).text = str(el)

        chem_table.cell(1, 0).text = "Min"
        chem_table.cell(2, 0).text = "Max"

        for i, el in enumerate(elements):
            chem_table.cell(1, i + 1).text = str(chem.get("min", {}).get(el, "-"))
            chem_table.cell(2, i + 1).text = str(chem.get("max", {}).get(el, "-"))
    else:
        doc.add_paragraph("Not specified")

    # -------- 4. MECHANICAL PROPERTIES --------
    doc.add_heading("4. MECHANICAL PROPERTIES", level=2)

    mech_props = data.get("mechanical_properties", [])
    if mech_props:
        mech_table = doc.add_table(rows=1, cols=4)
        mech_table.style = "Table Grid"

        hdr = mech_table.rows[0].cells
        hdr[0].text = "Property"
        hdr[1].text = "Min"
        hdr[2].text = "Max"
        hdr[3].text = "Unit"

        for m in mech_props:
            row = mech_table.add_row().cells
            row[0].text = str(m.get("property", "-"))
            row[1].text = str(m.get("min", "-"))
            row[2].text = str(m.get("max", "-"))
            row[3].text = str(m.get("unit", "-"))
    else:
        doc.add_paragraph("Not specified")

    # -------- 5. METALLURGICAL PROPERTIES --------
    doc.add_heading("5. METALLURGICAL PROPERTIES", level=2)

    mp = data.get("metallurgical_properties", {})
    doc.add_paragraph(f"Grain Size: {mp.get('grain_size', '-')}")
    doc.add_paragraph(f"Inclusion Rating: {mp.get('inclusion_rating', '-')}")
    doc.add_paragraph(f"Microstructure: {mp.get('microstructure', '-')}")
    doc.add_paragraph(f"Surface Condition: {mp.get('surface_condition', '-')}")
    doc.add_paragraph(f"Heat Treatment: {mp.get('heat_treatment', '-')}")

    # -------- 6. RAW TABLES FROM DOCUMENT --------
    if page_filtered_tables:
        doc.add_heading("6. RAW TABLES FROM SOURCE DOCUMENT", level=2)
        doc.add_paragraph(
            "The following tables were extracted directly from the source document. "
            "They are included here for reference and traceability."
        )

        for t_idx, t in enumerate(page_filtered_tables):
            rows = t.get("rows", [])

            if not rows or all(len(r) == 0 for r in rows):
                continue

            doc.add_paragraph(f"Table {t_idx + 1}  (Source Page: {t.get('page', 'Unknown')})")

            col_count = max(len(r) for r in rows)
            if col_count == 0:
                continue

            raw_table = doc.add_table(rows=len(rows), cols=col_count)
            raw_table.style = "Table Grid"

            for r_idx, row_cells in enumerate(rows):
                for c_idx in range(col_count):
                    cell_text = str(row_cells[c_idx]).strip() if c_idx < len(row_cells) else ""
                    raw_table.cell(r_idx, c_idx).text = cell_text

            doc.add_paragraph("") 

    # -------- FOOTER --------
    doc.add_paragraph(f"Source Pages: {sorted(list(selected_pages))}")

    doc.save(report_path)

    # -----------------------------
    # 7. RETURN FILE
    # -----------------------------
    return FileResponse(
        report_path,
        filename=f"doc_{doc_id}_Metallurgy_Report.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
# @app.get("/ingest/progress/{doc_id}")
# async def get_ingest_progress(doc_id: str):
#     return upload_progress.get(doc_id, {"status": "unknown", "percentage": 0}